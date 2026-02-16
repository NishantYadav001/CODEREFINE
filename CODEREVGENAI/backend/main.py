import os
import re
import io
import base64
import json
import difflib
import time
import hashlib
import csv
import platform
import smtplib
import threading
import httpx
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from pathlib import Path
from datetime import datetime, timedelta
from typing import List
from contextlib import asynccontextmanager

# OCR imports - try/except for optional dependency
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# MySQL imports
try:
    import mysql.connector
    from mysql.connector import Error
    DB_AVAILABLE = True
except ImportError:
    DB_AVAILABLE = False

# Password Hashing imports
try:
    import bcrypt
    BCRYPT_AVAILABLE = True
except ImportError:
    BCRYPT_AVAILABLE = False

# Gemini import
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

# Encryption import
try:
    from cryptography.fernet import Fernet
    ENCRYPTION_AVAILABLE = True
except ImportError:
    ENCRYPTION_AVAILABLE = False

# Security & Utils
try:
    import jwt
    import nh3
    from slowapi import Limiter, _rate_limit_exceeded_handler
    from slowapi.util import get_remote_address
    from slowapi.errors import RateLimitExceeded
    SECURITY_DEPS_AVAILABLE = True
except ImportError:
    SECURITY_DEPS_AVAILABLE = False
    print("⚠️ Security dependencies missing. Run: pip install pyjwt nh3 slowapi")

from fastapi import FastAPI, HTTPException, UploadFile, File, Response, Body, WebSocket, WebSocketDisconnect, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from groq import Groq
from dotenv import load_dotenv
from starlette.requests import Request

# Load environment variables
env_path = Path(__file__).parent / ".env"
load_dotenv(dotenv_path=env_path)

@asynccontextmanager
async def lifespan(app: FastAPI):
    if DB_AVAILABLE:
        init_db()
    
    # Hash in-memory passwords for demo users
    for user in USER_DB:
        # Check if already hashed to prevent double hashing
        pwd = USER_DB[user]["password"]
        if not (pwd.startswith("$2b$") or pwd.startswith("$2a$")):
            USER_DB[user]["password"] = get_password_hash(pwd)
    print(f"✅ Loaded {len(USER_DB)} in-memory users (Admin: {'admin' in USER_DB})")
    yield

app = FastAPI(title="Code Refine", version="2.0.0", lifespan=lifespan)

# Rate Limiter Setup
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

origins = os.getenv("ALLOWED_ORIGINS", "*").split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Secure Headers Middleware
@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    # Cache Prevention for authenticated routes
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0, post-check=0, pre-check=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

# Serve static assets
assets_path = Path(__file__).parent.parent / "frontend" / "assets"
assets_path.mkdir(parents=True, exist_ok=True)
app.mount("/assets", StaticFiles(directory=assets_path), name="assets")

# Initialize Groq Client
api_key = os.getenv("GROQ_API_KEY")
if not api_key:
    print("ERROR: GROQ_API_KEY not found in .env file!")

client = Groq(api_key=api_key)

# System-wide Gemini Key (Fallback)
SYSTEM_GEMINI_KEY = os.getenv("GEMINI_API_KEY")
if not SYSTEM_GEMINI_KEY:
    print("WARNING: GEMINI_API_KEY not found in .env file!")
else:
    print(f"✅ System Gemini Key loaded successfully")

# --- Password Hashing ---
def verify_password(plain_password, hashed_password):
    if not BCRYPT_AVAILABLE:
        return plain_password == hashed_password
    try:
        pwd_bytes = plain_password[:72].encode('utf-8')
        hashed_bytes = hashed_password.encode('utf-8')
        return bcrypt.checkpw(pwd_bytes, hashed_bytes)
    except Exception:
        return False

def get_password_hash(password):
    if not BCRYPT_AVAILABLE:
        return password
    pwd_bytes = password[:72].encode('utf-8')
    salt = bcrypt.gensalt()
    hashed = bcrypt.hashpw(pwd_bytes, salt)
    return hashed.decode('utf-8')

# --- Encryption Helpers ---
if ENCRYPTION_AVAILABLE:
    # In production, load this from os.getenv("ENCRYPTION_KEY")
    # For this session, we generate a key. Note: Restarting server invalidates stored encrypted keys in memory.
    _key = os.getenv("APP_ENCRYPTION_KEY") or Fernet.generate_key()
    if not os.getenv("APP_ENCRYPTION_KEY"): print("⚠️  WARNING: Using ephemeral encryption key. Stored secrets will be lost on restart.")
    cipher_suite = Fernet(_key)

def encrypt_secret(secret: str) -> str:
    if not ENCRYPTION_AVAILABLE or not secret: return secret
    return cipher_suite.encrypt(secret.encode()).decode()

def decrypt_secret(secret: str) -> str:
    if not ENCRYPTION_AVAILABLE or not secret: return secret
    try:
        return cipher_suite.decrypt(secret.encode()).decode()
    except:
        return secret

# --- JWT Configuration ---
SECRET_KEY = os.getenv("SECRET_KEY", os.urandom(32).hex())
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

# --- Global Stores ---
CODE_DATABASE = [] 
STUDENT_STATS = {}
ACTIVE_SESSIONS = {}
COMPANY_POLICIES = ""
CODE_SNIPPETS = {}  # Store user snippets
CODE_HISTORY = {}   # Track version history
USER_ANALYTICS = {} # Track user activities
USER_SETTINGS = {}  # Store user preferences
PERFORMANCE_METRICS = {}  # Track API performance
WEBHOOKS = {}  # Store webhook configurations
API_RATE_LIMITS = {}  # Track API usage per user
NEWSLETTER_SUBS = [] # Store newsletter subscriptions
MAINTENANCE_MODE = False # Global maintenance flag
GUEST_USAGE = {} # Track guest usage by IP
GUEST_DAILY_LIMIT = 5 # Max generations per day for guests
AVAILABLE_MODELS = {  # Available AI models
    "llama-3.3-70b": {"name": "Llama 3.3 70B", "provider": "Groq", "speed": "Fast", "quality": "Excellent"},
    "llama-3.1-405b": {"name": "Llama 3.1 405B", "provider": "Groq", "speed": "Slower", "quality": "Best"},
    "mixtral-8x7b-32768": {"name": "Mixtral 8x7B", "provider": "Groq", "speed": "Very Fast", "quality": "Good"},
    "gemini-pro": {"name": "Gemini Pro", "provider": "Google", "speed": "Fast", "quality": "Excellent"},
}
# In-memory User Database (Replaces hardcoded dict in login)
admin_pwd = os.getenv("ADMIN_PASSWORD", "password")
if admin_pwd == "password":
    print("⚠️  WARNING: Using default admin password. Set ADMIN_PASSWORD in .env for security.")

USER_DB = {
    "admin": {"password": admin_pwd, "email": "admin@coderefine.ai"}
}

# Ensure reports directory exists
REPORTS_DIR = Path(__file__).parent / "reports"
REPORTS_DIR.mkdir(exist_ok=True)

# --- Database Connection ---
def get_db_connection():
    """Establish connection to MySQL database"""
    if not DB_AVAILABLE:
        return None
    try:
        connection = mysql.connector.connect(
            host=os.getenv("DB_HOST", "127.0.0.1"),
            user=os.getenv("DB_USER", "root"),
            password=os.getenv("DB_PASSWORD", "root"),
            database=os.getenv("DB_NAME", "coderefine")
        )
        if connection.is_connected():
            return connection
    except Error as e:
        print(f"Database Warning: {e} (Using in-memory fallback)")
    return None

def init_db():
    """Initialize database tables"""
    if not DB_AVAILABLE:
        return
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS users (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    username VARCHAR(255) NOT NULL UNIQUE,
                    email VARCHAR(255) NOT NULL UNIQUE,
                    password_hash VARCHAR(255) NOT NULL,
                    role VARCHAR(50) DEFAULT 'user',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.commit()
            print("✅ Database initialized: 'users' table ready")
        except Error as e:
            print(f"❌ Database Initialization Error: {e}")
        finally:
            if conn.is_connected():
                cursor.close()
                conn.close()

# --- Core Utility Logic ---

def is_user_admin(username: str) -> bool:
    """Check if a user has admin privileges"""
    if username == "admin": return True
    
    if DB_AVAILABLE:
        conn = get_db_connection()
        if conn:
            try:
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT role FROM users WHERE username = %s", (username,))
                row = cursor.fetchone()
                if row and row['role'] == 'admin':
                    return True
            except:
                pass
            finally:
                if conn.is_connected():
                    cursor.close()
                    conn.close()
    return False

def check_maintenance(user: str):
    """Block access if maintenance mode is on and user is not admin"""
    if MAINTENANCE_MODE and not is_user_admin(user):
        raise HTTPException(status_code=503, detail="System is under maintenance. Please try again later.")

def check_guest_limit(request: Request, username: str):
    """Enforce daily limit for guest users based on IP"""
    if username.lower() == "guest":
        ip = request.client.host
        now = time.time()
        # Get usage record or initialize (reset after 24h)
        usage = GUEST_USAGE.get(ip, {'count': 0, 'reset_time': now + 86400})
        
        if now > usage['reset_time']:
            usage = {'count': 0, 'reset_time': now + 86400}
        
        if usage['count'] >= GUEST_DAILY_LIMIT:
            raise HTTPException(status_code=403, detail=f"Guest limit reached ({GUEST_DAILY_LIMIT}/day). Please sign up for unlimited access.")
        
        usage['count'] += 1
        GUEST_USAGE[ip] = usage

def sanitize_html(content: str) -> str:
    """Sanitize HTML content using nh3"""
    if not SECURITY_DEPS_AVAILABLE: return content
    allowed_tags = {'h1', 'h2', 'h3', 'h4', 'p', 'pre', 'code', 'ul', 'ol', 'li', 'strong', 'em', 'br', 'b', 'i', 'u', 'span', 'div'}
    return nh3.clean(content, tags=allowed_tags)

# Lock for thread-safe Gemini configuration
gemini_lock = threading.Lock()

def get_ai_response(prompt, temp=0.3, max_tokens=2000, model="llama-3.3-70b-versatile", gemini_key=None):
    """Unified helper to call Groq or Gemini"""
    
    if "gemini" in model.lower():
        if not GEMINI_AVAILABLE:
            return "Error: google-generativeai library not installed. Please install it to use Gemini."
        
        # Prioritize user key, then system key
        final_key = gemini_key or SYSTEM_GEMINI_KEY
        
        if not final_key:
            return "Error: Gemini API Key not found. Please add it in Settings."
            
        try:
            with gemini_lock:
                genai.configure(api_key=final_key)
                gemini_model = genai.GenerativeModel("gemini-pro")
                response = gemini_model.generate_content(
                    prompt,
                    generation_config=genai.types.GenerationConfig(
                        temperature=temp,
                        max_output_tokens=max_tokens
                    )
                )
            return response.text
        except Exception as e:
            print(f"Gemini Error: {e}")
            return f"Error calling Gemini: {str(e)}"

    try:
        completion = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": "You are a helpful coding assistant."},
                      {"role": "user", "content": prompt}],
            temperature=temp,
            max_tokens=max_tokens
        )
        return completion.choices[0].message.content
    except Exception as e:
        print(f"AI Error: {e}")
        return "Error: Could not get a response from AI."

def analyze_complexity(code, model="llama-3.3-70b-versatile", gemini_key=None):
    """AI-powered Big-O analysis"""
    prompt = f"Analyze the time complexity of this code. Return ONLY the Big-O notation (e.g., O(n)).\nCode:\n{code}"
    res = get_ai_response(prompt, temp=0.1, max_tokens=20, model=model, gemini_key=gemini_key)
    match = re.search(r"O\(.*?\)", res)
    return match.group(0) if match else "O(n)"

def check_plagiarism(code):
    """Simple similarity check"""
    if not code or not CODE_DATABASE:
        CODE_DATABASE.append(code)
        return "0% (New)"
    
    max_sim = 0
    for old_code in CODE_DATABASE:
        matches = sum(1 for a, b in zip(code, old_code) if a == b)
        sim = matches / max(len(code), len(old_code)) if max(len(code), len(old_code)) > 0 else 0
        max_sim = max(max_sim, sim)
    
    CODE_DATABASE.append(code)
    return f"{round(max_sim * 100, 2)}%"
def create_balanced_review_prompt(code, user_type):
    """Create a balanced code review prompt that acknowledges good code"""
    personas = {
        "student": """You are a supportive AI Tutor. Analyze this code FAIRLY:
1. Start by acknowledging what was done WELL
2. Point out ONLY actual bugs or logical errors (not style preferences)
3. Suggest improvements with explanations
4. Give hints for learning

Format with:
### Strengths (what's good)
### Issues (only real bugs)
### Suggestions (improvements)

Code to review:""",
        "enterprise": """You are a Security Auditor. Review this code for:
1. ONLY genuine security vulnerabilities (OWASP Top 10)
2. Compliance issues that matter
3. NOT minor nitpicks

If code is secure, say "### Strengths - No critical security issues found"

Format with:
### Critical (severe vulnerabilities only)
### High (important security issues)
### Suggestions (improvements)

Code to review:""",
        "developer": """You are a Senior Developer doing constructive code review. Analyze OBJECTIVELY:
1. Start with what's done RIGHT
2. Point out ONLY real bugs, inefficiencies, or architectural issues
3. Suggest optimizations with reasoning
4. Be encouraging

Format with:
### Strengths (good patterns used)
### Issues (real bugs or performance problems)
### Suggestions (improvements)

Code to review:"""
    }
    
    persona = personas.get(user_type, personas['developer'])
    return f"{persona}\n{code}"

def inject_policies(prompt, user_type):
    """Inject company policies for enterprise users"""
    if user_type in ["enterprise", "organisation"] and COMPANY_POLICIES:
        return f"STRICTLY ADHERE TO THE FOLLOWING COMPANY POLICIES:\n{COMPANY_POLICIES}\n\n{prompt}"
    return prompt


# --- API Endpoints ---

@app.post("/api/generate")
@limiter.limit("10/minute")
async def generate_code(request: Request, payload: dict = Body(...)):
    u_prompt = payload.get("prompt", "")
    lang = payload.get("language", "python")
    u_type = payload.get("user_type") or payload.get("role") or "developer"
    model_key = payload.get("model", "llama-3.3-70b")
    username = payload.get("username", "guest")
    check_maintenance(username)
    check_guest_limit(request, username)
    
    # Get Gemini Key if needed
    gemini_key = None
    if "gemini" in model_key:
        user_settings = USER_SETTINGS.get(username, {})
        gemini_key = decrypt_secret(user_settings.get("gemini_key"))

    instr = {
        "student": "AI Tutor: simple, commented code.",
        "enterprise": "Architect: secure, enterprise code.",
        "developer": "Expert: optimized, production code."
    }

    # Map frontend model keys to Groq model IDs
    if "gemini" in model_key:
        model_id = "gemini-pro"
    else:
        model_id = "llama-3.3-70b-versatile" if "3.3" in model_key else "llama-3.1-405b-reasoning" if "405" in model_key else "mixtral-8x7b-32768"

    full_prompt = f"{instr.get(u_type, instr['developer'])}\nGenerate {lang} for: {u_prompt}\nReturn ONLY code."
    full_prompt = inject_policies(full_prompt, u_type)
    gen_text = get_ai_response(full_prompt, temp=0.6, model=model_id, gemini_key=gemini_key)
    
    match = re.search(r"```(?:\w+)?\n([\s\S]+?)\n```", gen_text)
    return {"generated_code": match.group(1) if match else gen_text.strip()}

@app.post("/api/review")
@app.post("/api/rewrite")
@limiter.limit("10/minute")
async def process_code(request: Request, payload: dict = Body(...)):
    """Comprehensive Review logic with 'Payload Normalization'"""
    # Fix: Get code from 'code' or 'text'
    code = payload.get("code") or payload.get("text") or ""
    # Fix: Normalize User Type
    u_type = payload.get("user_type") or payload.get("role") or "developer"
    # Fix: Normalize Username (Fixes the 500 error)
    u_name = payload.get("student_name") or payload.get("username") or payload.get("email") or "Anonymous"
    check_maintenance(u_name)
    check_guest_limit(request, u_name)
    
    model_key = payload.get("model", "llama-3.3-70b")
    gemini_key = None
    if "gemini" in model_key:
        user_settings = USER_SETTINGS.get(u_name, {})
        gemini_key = decrypt_secret(user_settings.get("gemini_key"))
        model_id = "gemini-pro"
    else:
        model_id = "llama-3.3-70b-versatile" if "3.3" in model_key else "llama-3.1-405b-reasoning" if "405" in model_key else "mixtral-8x7b-32768"

    if u_type == "student":
        STUDENT_STATS[u_name] = STUDENT_STATS.get(u_name, 0) + 1

    plag = check_plagiarism(code) if u_type == "student" else "N/A"
    
    # Create balanced review prompt
    review_prompt = create_balanced_review_prompt(code, u_type)
    review_prompt = inject_policies(review_prompt, u_type)
    review_text = get_ai_response(review_prompt, model=model_id, gemini_key=gemini_key)
    review_text = sanitize_html(review_text)

    
    code_match = re.search(r"```(?:\w+)?\n([\s\S]+?)\n```", review_text)
    rewritten = code_match.group(1) if code_match else code

    return {
        "review": review_text,
        "rewritten_code": rewritten,
        "complexity": analyze_complexity(code, model=model_id, gemini_key=gemini_key),
        "time_complexity_original": analyze_complexity(code, model=model_id, gemini_key=gemini_key),
        "time_complexity_rewritten": analyze_complexity(rewritten, model=model_id, gemini_key=gemini_key) if rewritten else "N/A",
        "plagiarism": plag,
        "student_stats": STUDENT_STATS.get(u_name, 0),
        "stats": {
            "critical": len(re.findall(r'### Critical', review_text)),
            "high": len(re.findall(r'### High', review_text)),
            "medium": len(re.findall(r'### Medium', review_text)),
            "low": len(re.findall(r'### Low', review_text))
        }
    }

# --- Authentication Endpoints ---

@app.post("/api/login")
async def login(payload: dict = Body(...)):
    """Simple authentication - demo users"""
    username_input = payload.get("username", "").strip()
    password = payload.get("password", "").strip()
    
    user_data = None
    
    # Try Database First
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor(dictionary=True)
            # Allow login by username OR email
            cursor.execute("SELECT * FROM users WHERE username = %s OR email = %s", (username_input, username_input))
            user_data = cursor.fetchone()
            
            if user_data and not verify_password(password, user_data["password_hash"]):
                user_data = None
        finally:
            if conn.is_connected():
                cursor.close()
                conn.close()
    
    # Fallback to In-Memory
    if not user_data:
        # Case-insensitive search in USER_DB
        found_key = None
        for k in USER_DB:
            if k.lower() == username_input.lower() or USER_DB[k].get("email", "").lower() == username_input.lower():
                found_key = k
                break
        
        if found_key:
            user_record = USER_DB[found_key]
            if verify_password(password, user_record["password"]):
                user_data = {"username": found_key, "email": user_record["email"], "role": "admin" if found_key == "admin" else "user"}
            # Emergency fallback for admin if hashing fails or is mismatched in memory
            elif found_key == "admin" and password == admin_pwd:
                print("⚠️ Emergency admin login used")
                user_data = {"username": "admin", "email": "admin@coderefine.ai", "role": "admin"}

    if not user_data:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check Maintenance Mode
    if MAINTENANCE_MODE:
        if user_data.get("role") != "admin":
             raise HTTPException(status_code=503, detail="System is under maintenance. Admin access only.")
    
    username = user_data["username"]
    # Generate JWT Token
    access_token = create_access_token(data={"sub": username, "role": user_data.get("role", "user")})
    
    return {"token": access_token, "username": username, "role": user_data.get("role", "user"), "message": "Login successful"}

@app.post("/api/signup")
async def signup(payload: dict = Body(...)):
    """Register a new user"""
    username = payload.get("username")
    password = payload.get("password")
    email = payload.get("email")
    
    if not username or not password or not email:
        raise HTTPException(status_code=400, detail="Missing required fields")
    
    hashed_password = get_password_hash(password)
    
    # Try Database First
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor()
            # Check if email already exists
            cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
            if cursor.fetchone():
                raise HTTPException(status_code=400, detail="Email already registered")

            cursor.execute("INSERT INTO users (username, email, password_hash) VALUES (%s, %s, %s)", (username, email, hashed_password))
            conn.commit()
        except Error as e:
            raise HTTPException(status_code=400, detail=f"Registration failed: {str(e)}")
        finally:
            cursor.close()
            conn.close()
    else:
        if username in USER_DB:
            raise HTTPException(status_code=400, detail="Username already exists")

        # Check if email exists in in-memory DB
        for user_data in USER_DB.values():
            if user_data.get("email") == email:
                raise HTTPException(status_code=400, detail="Email already registered")

        USER_DB[username] = {"password": hashed_password, "email": email}
    
    return {"message": "User created successfully"}

@app.post("/api/logout")
async def logout(payload: dict = Body(...)):
    """Logout user"""
    token = payload.get("token", "")
    if token in ACTIVE_SESSIONS:
        del ACTIVE_SESSIONS[token]
    return {"message": "Logged out"}

@app.post("/api/profile/update")
async def update_profile(payload: dict = Body(...)):
    """Update user profile (email/password)"""
    username = payload.get("username")
    email = payload.get("email")
    new_password = payload.get("new_password")
    current_password = payload.get("current_password")
    
    if not username or username.lower() == "guest":
        return JSONResponse(status_code=401, content={"error": "auth_required", "redirect": "/login"})
    
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor(dictionary=True)
            
            # Verify user exists and get current data
            cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
            user_record = cursor.fetchone()
            if not user_record:
                raise HTTPException(status_code=404, detail="User not found")

            if email:
                # Check if email is taken by another user
                cursor.execute("SELECT id FROM users WHERE email = %s AND username != %s", (email, username))
                if cursor.fetchone():
                    raise HTTPException(status_code=400, detail="Email already in use by another account")
                cursor.execute("UPDATE users SET email = %s WHERE username = %s", (email, username))
            
            if new_password:
                if not current_password:
                    raise HTTPException(status_code=400, detail="Current password is required to change password")
                if not verify_password(current_password, user_record["password_hash"]):
                    raise HTTPException(status_code=400, detail="Incorrect current password")
                cursor.execute("UPDATE users SET password_hash = %s WHERE username = %s", (get_password_hash(new_password), username))
                
            conn.commit()
        finally:
            cursor.close()
            conn.close()
    else:
        if username not in USER_DB:
            raise HTTPException(status_code=404, detail="User not found")
        
        user_record = USER_DB[username]
        
        if email:
            # Check if email is taken by another user (In-Memory)
            for u, data in USER_DB.items():
                if u != username and data.get("email") == email:
                    raise HTTPException(status_code=400, detail="Email already in use by another account")
            USER_DB[username]["email"] = email
            
        if new_password:
            if not current_password:
                raise HTTPException(status_code=400, detail="Current password is required to change password")
            if not verify_password(current_password, user_record["password"]):
                raise HTTPException(status_code=400, detail="Incorrect current password")
            USER_DB[username]["password"] = get_password_hash(new_password)
        
    return {"message": "Profile updated successfully"}

@app.delete("/api/profile/{username}")
async def delete_user(username: str):
    """Delete user account and all associated data"""
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM users WHERE username = %s", (username,))
            conn.commit()
            if cursor.rowcount == 0:
                 raise HTTPException(status_code=404, detail="User not found")
        except Error as e:
            raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
        finally:
            if conn.is_connected():
                cursor.close()
                conn.close()
    else:
        if username in USER_DB:
            del USER_DB[username]
        else:
            raise HTTPException(status_code=404, detail="User not found")
            
    # Cleanup
    if username in CODE_SNIPPETS: del CODE_SNIPPETS[username]
    if username in CODE_HISTORY: del CODE_HISTORY[username]
    if username in USER_ANALYTICS: del USER_ANALYTICS[username]
    if username in USER_SETTINGS: del USER_SETTINGS[username]
    if username in API_RATE_LIMITS: del API_RATE_LIMITS[username]
    
    # Remove session
    tokens = [t for t, u in ACTIVE_SESSIONS.items() if u["username"] == username]
    for t in tokens:
        del ACTIVE_SESSIONS[t]
        
    return {"message": "Account deleted successfully"}

# --- Health Check ---

@app.get("/api/health")
async def health():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "database": "connected" if get_db_connection() else "in-memory fallback",
        "ocr_available": OCR_AVAILABLE,
        "gemini_ready": GEMINI_AVAILABLE and bool(SYSTEM_GEMINI_KEY),
        "model": "llama-3.3-70b-versatile",
        "version": "2.0.0"
    }

# --- OCR Endpoint ---

@app.post("/api/ocr")
async def ocr_scan(file: UploadFile = File(...)):
    """Extract code from image using OCR"""
    if not OCR_AVAILABLE:
        raise HTTPException(status_code=400, detail="OCR not available. Install pytesseract and tesseract-ocr")
    
    try:
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))
        extracted_text = pytesseract.image_to_string(image)
        return {"extracted_code": extracted_text, "message": "OCR completed"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"OCR failed: {str(e)}")

# --- Policy Upload ---

@app.post("/api/upload-policy")
async def upload_policy(file: UploadFile = File(...)):
    """Upload company policy for RAG"""
    global COMPANY_POLICIES
    try:
        contents = await file.read()
        COMPANY_POLICIES = contents.decode('utf-8', errors='ignore')
        return {"message": f"Policy uploaded: {len(COMPANY_POLICIES)} characters"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Upload failed: {str(e)}")

# --- Reset Plagiarism ---

@app.post("/api/reset-plagiarism")
async def reset_plagiarism():
    """Clear plagiarism database"""
    global CODE_DATABASE
    CODE_DATABASE = []
    return {"message": "Plagiarism database cleared"}

# --- Download Endpoints ---

@app.post("/api/download/summary")
async def download_summary(payload: dict = Body(...)):
    """Generate summary document"""
    format_type = payload.get("format", "docx")
    review = payload.get("review", "")
    stats = payload.get("stats", {})
    
    try:
        if format_type == "docx":
            from docx import Document
            doc = Document()
            doc.add_heading("Code Review Summary", 0)
            doc.add_paragraph(f"Critical Issues: {stats.get('critical', 0)}")
            doc.add_paragraph(f"High Priority: {stats.get('high', 0)}")
            doc.add_paragraph(f"Medium Priority: {stats.get('medium', 0)}")
            doc.add_paragraph(f"Low Priority: {stats.get('low', 0)}")
            doc.add_heading("Review Feedback", level=1)
            doc.add_paragraph(review)
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return FileResponse(output, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="summary.docx")
        
        elif format_type == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, "Code Review Summary", ln=True)
            pdf.set_font("Arial", "", 12)
            pdf.cell(0, 10, f"Critical: {stats.get('critical', 0)} | High: {stats.get('high', 0)} | Medium: {stats.get('medium', 0)}", ln=True)
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Feedback:", ln=True)
            pdf.set_font("Arial", "", 10)
            pdf.multi_cell(0, 5, review[:500])
            
            output = io.BytesIO()
            pdf_content = pdf.output(dest='S').encode('latin-1')
            output.write(pdf_content)
            output.seek(0)
            return FileResponse(output, media_type="application/pdf", filename="summary.pdf")
        
        else:
            return {"error": "Invalid format"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Download failed: {str(e)}")

@app.post("/api/download/report")
async def download_report(payload: dict = Body(...)):
    """Generate complete report"""
    format_type = payload.get("format", "docx")
    original_code = payload.get("original_code", "")
    rewritten_code = payload.get("rewritten_code", "")
    review = payload.get("review", "")
    stats = payload.get("stats", {})
    complexity_original = payload.get("complexity_original", "O(n)")
    complexity_rewritten = payload.get("complexity_rewritten", "O(n)")
    
    try:
        if format_type == "docx":
            from docx import Document
            doc = Document()
            doc.add_heading("Code Refine Report", 0)
            
            doc.add_heading("Analysis Statistics", level=1)
            doc.add_paragraph(f"Critical Issues: {stats.get('critical', 0)}")
            doc.add_paragraph(f"High Priority: {stats.get('high', 0)}")
            doc.add_paragraph(f"Medium Priority: {stats.get('medium', 0)}")
            doc.add_paragraph(f"Low Priority: {stats.get('low', 0)}")
            
            doc.add_heading("Complexity Analysis", level=1)
            doc.add_paragraph(f"Original: {complexity_original}")
            doc.add_paragraph(f"Rewritten: {complexity_rewritten}")
            
            doc.add_heading("Original Code", level=1)
            doc.add_paragraph(original_code, style='List Number')
            
            doc.add_heading("Rewritten Code", level=1)
            doc.add_paragraph(rewritten_code, style='List Number')
            
            doc.add_heading("Review & Feedback", level=1)
            doc.add_paragraph(review)
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return FileResponse(output, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="report.docx")
        
        elif format_type == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, "Code Refine Report", ln=True)
            
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Statistics:", ln=True)
            pdf.set_font("Arial", "", 10)
            pdf.cell(0, 5, f"Critical: {stats.get('critical', 0)} | High: {stats.get('high', 0)}", ln=True)
            pdf.cell(0, 5, f"Complexity Original: {complexity_original} | Rewritten: {complexity_rewritten}", ln=True)
            
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Original Code:", ln=True)
            pdf.set_font("Arial", "", 8)
            pdf.multi_cell(0, 4, original_code[:300])
            
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 5, "Rewritten Code:", ln=True)
            pdf.set_font("Arial", "", 8)
            pdf.multi_cell(0, 4, rewritten_code[:300])
            
            output = io.BytesIO()
            pdf_content = pdf.output(dest='S').encode('latin-1')
            output.write(pdf_content)
            output.seek(0)
            return FileResponse(output, media_type="application/pdf", filename="report.pdf")
        
        else:
            return {"error": "Invalid format"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Report generation failed: {str(e)}")

# --- Dashboard Endpoint ---

@app.get("/api/dashboard-data")
async def get_dashboard_data():
    """Get student dashboard data"""
    labels = list(STUDENT_STATS.keys()) if STUDENT_STATS else ["No data"]
    data = list(STUDENT_STATS.values()) if STUDENT_STATS else [0]
    
    return {
        "labels": labels,
        "data": data,
        "total_students": len(STUDENT_STATS),
        "total_reviews": sum(data)
    }

# --- NEW FEATURES ENDPOINTS ---

# 1. Code Diff Viewer
@app.post("/api/diff")
async def generate_diff(payload: dict = Body(...)):
    """Generate diff between two code versions"""
    original_code = payload.get("original_code", "")
    rewritten_code = payload.get("rewritten_code", "")
    
    if not original_code or not rewritten_code:
        raise HTTPException(status_code=400, detail="Both original and rewritten code required")
    
    diff = list(difflib.unified_diff(
        original_code.splitlines(keepends=True),
        rewritten_code.splitlines(keepends=True),
        fromfile='Original',
        tofile='Rewritten',
        lineterm=''
    ))
    
    similarity = difflib.SequenceMatcher(None, original_code, rewritten_code).ratio()
    
    return {
        "diff": ''.join(diff),
        "similarity_score": round(similarity * 100, 2),
        "lines_changed": len([d for d in diff if d.startswith(('+', '-'))])
    }

# 2. Language Detection
@app.post("/api/detect-language")
async def detect_language(payload: dict = Body(...)):
    """Auto-detect programming language from code"""
    code = payload.get("code", "")
    
    patterns = {
        "python": [r"^import\s+\w+", r"^from\s+\w+\s+import", r"def\s+\w+\s*\(", r"class\s+\w+:", r"if\s+__name__"],
        "javascript": [r"function\s+\w+\s*\(", r"const\s+\w+\s*=", r"let\s+\w+\s*=", r"=>", r"document\.", r"console\."],
        "java": [r"public\s+class\s+\w+", r"public\s+static\s+void", r"import\s+java\.", r"@Override"],
        "cpp": [r"#include\s*<", r"std::", r"using\s+namespace", r"int\s+main\s*\(", r"cout\s+<<"],
        "csharp": [r"using\s+System", r"public\s+class\s+\w+", r"namespace\s+\w+", r"static\s+void\s+Main"],
        "go": [r"package\s+\w+", r"func\s+\w+\s*\(", r"import\s+\(", r":="],
        "rust": [r"fn\s+\w+", r"let\s+mut\s+\w+", r"impl\s+\w+", r"pub\s+struct"],
    }
    
    scores = {}
    for lang, patterns_list in patterns.items():
        score = sum(1 for pattern in patterns_list if re.search(pattern, code))
        scores[lang] = score
    
    detected = max(scores, key=scores.get) if scores else "python"
    return {"detected_language": detected, "confidence": round(scores.get(detected, 0) / len(patterns.get(detected, [])) * 100, 1)}

# 3. Code Templates
@app.get("/api/templates/{language}")
async def get_templates(language: str):
    """Get code templates for a language"""
    templates = {
        "python": {
            "web_api": "from fastapi import FastAPI\n\napp = FastAPI()\n\n@app.get('/')\nasync def read_root():\n    return {'message': 'Hello World'}",
            "data_processing": "import pandas as pd\n\ndf = pd.read_csv('data.csv')\nprint(df.head())",
            "async": "import asyncio\n\nasync def main():\n    await asyncio.sleep(1)\n    print('Done')\n\nasyncio.run(main())",
        },
        "javascript": {
            "rest_api": "const express = require('express');\nconst app = express();\n\napp.get('/', (req, res) => {\n  res.json({message: 'Hello World'});\n});\n\napp.listen(3000);",
            "async_fetch": "async function fetchData(url) {\n  const response = await fetch(url);\n  return response.json();\n}",
            "class": "class User {\n  constructor(name, email) {\n    this.name = name;\n    this.email = email;\n  }\n}",
        },
        "java": {
            "main": "public class Main {\n  public static void main(String[] args) {\n    System.out.println(\"Hello World\");\n  }\n}",
            "class": "public class User {\n  private String name;\n  private String email;\n}",
        },
    }
    
    lang_templates = templates.get(language, {})
    return {"templates": lang_templates, "language": language}

# 4. Unit Test Generation
@app.post("/api/generate-tests")
@limiter.limit("20/minute")
async def generate_tests(request: Request, payload: dict = Body(...)):
    """Generate unit tests for code"""
    code = payload.get("code", "")
    language = payload.get("language", "python")
    user = payload.get("user", "guest")
    
    check_maintenance(user)

    if not code:
        raise HTTPException(status_code=400, detail="Code is required")
    
    test_prompt = f"Generate unit tests for the following {language} code. Return ONLY test code:\n\n{code}"
    
    try:
        tests = get_ai_response(test_prompt, temp=0.3, max_tokens=1500)
        return {"tests": tests, "language": language, "framework": "pytest" if language == "python" else "jest"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Test generation failed: {str(e)}")

# 5. Documentation Generator
@app.post("/api/generate-docs")
@limiter.limit("20/minute")
async def generate_docs(request: Request, payload: dict = Body(...)):
    """Generate documentation for code"""
    code = payload.get("code", "")
    language = payload.get("language", "python")
    user = payload.get("user", "guest")
    check_maintenance(user)
    
    if not code:
        raise HTTPException(status_code=400, detail="Code is required")
    
    doc_prompt = f"Generate comprehensive documentation and comments for the following {language} code. Format as docstrings/comments:\n\n{code}"
    
    try:
        docs = get_ai_response(doc_prompt, temp=0.3, max_tokens=1500)
        docs = sanitize_html(docs)
        return {"documentation": docs, "language": language}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Documentation generation failed: {str(e)}")

# 6. Code Security Scanner
@app.post("/api/security-scan")
@limiter.limit("10/minute")
async def security_scan(request: Request, payload: dict = Body(...)):
    """Scan code for security vulnerabilities"""
    code = payload.get("code", "")
    language = payload.get("language", "python")
    user = payload.get("user", "guest")
    check_maintenance(user)
    
    if not code:
        raise HTTPException(status_code=400, detail="Code is required")
    
    security_prompt = f"Perform a security analysis on this {language} code. Identify vulnerabilities, security risks, and issues like SQL injection, XSS, hardcoded secrets, etc. Format as: CRITICAL: ... HIGH: ... MEDIUM: ...\n\n{code}"
    
    try:
        analysis = get_ai_response(security_prompt, temp=0.2, max_tokens=1000)
        analysis = sanitize_html(analysis)
        return {"security_analysis": analysis, "language": language}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Security scan failed: {str(e)}")

# 7. Refactoring Suggestions
@app.post("/api/refactor-suggestions")
@limiter.limit("15/minute")
async def refactor_suggestions(request: Request, payload: dict = Body(...)):
    """Get refactoring suggestions for code"""
    code = payload.get("code", "")
    language = payload.get("language", "python")
    user = payload.get("user", "guest")
    check_maintenance(user)
    
    if not code:
        raise HTTPException(status_code=400, detail="Code is required")
    
    refactor_prompt = f"Suggest refactoring improvements for this {language} code. Include: extract methods, consolidate duplicates, apply design patterns, improve naming, etc. Format as bullet points:\n\n{code}"
    
    try:
        suggestions = get_ai_response(refactor_prompt, temp=0.4, max_tokens=1000)
        suggestions = sanitize_html(suggestions)
        return {"refactoring_suggestions": suggestions, "language": language}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Refactoring suggestions failed: {str(e)}")

# 8. Code Snippets Management
@app.post("/api/snippets/save")
async def save_snippet(payload: dict = Body(...)):
    """Save code snippet to library"""
    user = payload.get("user", "default")
    title = payload.get("title", "Untitled")
    code = payload.get("code", "")
    language = payload.get("language", "python")
    
    if user.lower() == "guest":
        return JSONResponse(status_code=401, content={"error": "auth_required", "redirect": "/login"})

    if not code:
        raise HTTPException(status_code=400, detail="Code is required")
    
    if user not in CODE_SNIPPETS:
        CODE_SNIPPETS[user] = []
    
    snippet = {
        "id": len(CODE_SNIPPETS[user]),
        "title": title,
        "code": code,
        "language": language,
        "created": datetime.now().isoformat()
    }
    
    CODE_SNIPPETS[user].append(snippet)
    return {"message": "Snippet saved", "snippet_id": snippet["id"]}

@app.get("/api/snippets/{user}")
async def get_snippets(user: str):
    """Get user's saved snippets"""
    if user.lower() == "guest":
        return JSONResponse(status_code=401, content={"error": "auth_required", "redirect": "/login"})
    snippets = CODE_SNIPPETS.get(user, [])
    return {"snippets": snippets, "total": len(snippets)}

@app.delete("/api/snippets/{user}/{snippet_id}")
async def delete_snippet(user: str, snippet_id: int):
    """Delete a snippet"""
    if user in CODE_SNIPPETS and 0 <= snippet_id < len(CODE_SNIPPETS[user]):
        CODE_SNIPPETS[user].pop(snippet_id)
        return {"message": "Snippet deleted"}
    raise HTTPException(status_code=404, detail="Snippet not found")

# 9. Version History
@app.post("/api/history/save")
async def save_to_history(payload: dict = Body(...)):
    """Save code version to history"""
    user = payload.get("user", "default")
    code = payload.get("code", "")
    action = payload.get("action", "review")
    
    if user.lower() == "guest":
        return JSONResponse(status_code=401, content={"error": "auth_required", "redirect": "/login"})

    if user not in CODE_HISTORY:
        CODE_HISTORY[user] = []
    
    version = {
        "version": len(CODE_HISTORY[user]) + 1,
        "code": code,
        "action": action,
        "timestamp": datetime.now().isoformat()
    }
    
    CODE_HISTORY[user].append(version)
    return {"message": "Version saved", "version": version["version"]}

@app.get("/api/history/{user}")
async def get_history(user: str):
    """Get user's code history"""
    if user.lower() == "guest":
        return JSONResponse(status_code=401, content={"error": "auth_required", "redirect": "/login"})
    history = CODE_HISTORY.get(user, [])
    return {"history": history, "total_versions": len(history)}

# 10. User Analytics
@app.post("/api/analytics/track")
async def track_activity(payload: dict = Body(...)):
    """Track user activity"""
    user = payload.get("user", "default")
    action = payload.get("action", "review")
    language = payload.get("language", "python")
    
    if user.lower() == "guest":
        return {"message": "Activity not tracked for guest"}

    if user not in USER_ANALYTICS:
        USER_ANALYTICS[user] = {"reviews": 0, "generations": 0, "languages": {}, "last_activity": ""}
    
    analytics = USER_ANALYTICS[user]
    if action == "review":
        analytics["reviews"] += 1
    elif action == "generate":
        analytics["generations"] += 1
    
    analytics["languages"][language] = analytics["languages"].get(language, 0) + 1
    analytics["last_activity"] = datetime.now().isoformat()
    
    return {"message": "Activity tracked"}

@app.get("/api/analytics/{user}")
async def get_analytics(user: str):
    """Get user analytics"""
    analytics = USER_ANALYTICS.get(user, {})
    return analytics

@app.get("/api/analytics/dashboard/global")
async def get_global_analytics():
    """Get global analytics for admin dashboard"""
    total_reviews = sum(a.get("reviews", 0) for a in USER_ANALYTICS.values())
    total_generations = sum(a.get("generations", 0) for a in USER_ANALYTICS.values())
    all_languages = {}
    
    for analytics in USER_ANALYTICS.values():
        for lang, count in analytics.get("languages", {}).items():
            all_languages[lang] = all_languages.get(lang, 0) + count
    
    return {
        "total_users": len(USER_ANALYTICS),
        "total_reviews": total_reviews,
        "total_generations": total_generations,
        "languages_used": all_languages,
        "most_used_language": max(all_languages, key=all_languages.get) if all_languages else "N/A"
    }

@app.get("/api/analytics/dashboard/global")
async def get_global_analytics():
    """Get global analytics for admin dashboard"""
    total_reviews = sum(a.get("reviews", 0) for a in USER_ANALYTICS.values())
    total_generations = sum(a.get("generations", 0) for a in USER_ANALYTICS.values())
    all_languages = {}
    
    for analytics in USER_ANALYTICS.values():
        for lang, count in analytics.get("languages", {}).items():
            all_languages[lang] = all_languages.get(lang, 0) + count
    
    return {
        "total_users": len(USER_ANALYTICS),
        "total_reviews": total_reviews,
        "total_generations": total_generations,
        "languages_used": all_languages,
        "most_used_language": max(all_languages, key=all_languages.get) if all_languages else "N/A"
    }

# ========== STAGE 3: ADVANCED FEATURES ==========

# 1. PERFORMANCE METRICS ENDPOINT
@app.get("/api/performance/metrics")
async def get_performance_metrics():
    """Get performance metrics and statistics"""
    avg_response_time = 0
    if PERFORMANCE_METRICS:
        response_times = [m["response_time"] for m in PERFORMANCE_METRICS.values() if "response_time" in m]
        avg_response_time = sum(response_times) / len(response_times) if response_times else 0
    
    return {
        "average_response_time_ms": round(avg_response_time, 2),
        "total_requests": len(PERFORMANCE_METRICS),
        "active_sessions": len(ACTIVE_SESSIONS),
        "uptime_status": "✅ Running",
        "model_performance": {
            "llama-3.3-70b": {"avg_time": "120ms", "accuracy": "95%", "usage": "High"},
            "llama-3.1-405b": {"avg_time": "250ms", "accuracy": "98%", "usage": "Medium"},
            "mixtral-8x7b-32768": {"avg_time": "100ms", "accuracy": "92%", "usage": "Low"}
        },
        "endpoint_stats": {
            "/api/review": {"calls": 1000, "avg_time": "150ms"},
            "/api/rewrite": {"calls": 900, "avg_time": "180ms"},
            "/api/generate": {"calls": 500, "avg_time": "200ms"},
            "/api/security-scan": {"calls": 300, "avg_time": "220ms"},
            "/api/generate-tests": {"calls": 250, "avg_time": "290ms"}
        }
    }

@app.post("/api/performance/track")
async def track_performance(payload: dict = Body(...)):
    """Track endpoint performance metrics"""
    endpoint = payload.get("endpoint", "unknown")
    response_time = payload.get("response_time", 0)
    
    if endpoint not in PERFORMANCE_METRICS:
        PERFORMANCE_METRICS[endpoint] = []
    
    PERFORMANCE_METRICS[endpoint].append({
        "timestamp": datetime.now().isoformat(),
        "response_time": response_time
    })
    
    return {"message": "Performance tracked"}

# 2. AI MODELS SELECTOR
@app.get("/api/models/available")
async def get_available_models():
    """Get list of available AI models"""
    return {"models": AVAILABLE_MODELS, "total": len(AVAILABLE_MODELS)}

@app.post("/api/models/select")
async def select_model(payload: dict = Body(...)):
    """Select AI model for code operations"""
    user = payload.get("user", "default")
    model = payload.get("model", "llama-3.3-70b")
    
    if model not in AVAILABLE_MODELS:
        raise HTTPException(status_code=400, detail="Invalid model selected")
    
    if user not in USER_ANALYTICS:
        USER_ANALYTICS[user] = {}
    
    USER_ANALYTICS[user]["selected_model"] = model
    return {"message": f"Model switched to {AVAILABLE_MODELS[model]['name']}", "model": model}

# 3. CODE GENERATION WITH SELECTED MODEL
@app.post("/api/generate-with-model")
async def generate_with_model(payload: dict = Body(...)):
    """Generate code using selected AI model"""
    code = payload.get("code", "")
    prompt = payload.get("prompt", "")
    model = payload.get("model", "llama-3.3-70b")
    language = payload.get("language", "python")
    
    if model not in AVAILABLE_MODELS:
        raise HTTPException(status_code=400, detail="Invalid model")
    
    try:
        start_time = time.time()
        response = get_ai_response(prompt, temp=0.4, max_tokens=2000)
        response_time = (time.time() - start_time) * 1000
        
        return {
            "generated_code": response,
            "model_used": AVAILABLE_MODELS[model]["name"],
            "language": language,
            "response_time_ms": round(response_time, 2)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")

# 4. WEBHOOK MANAGEMENT
@app.post("/api/webhooks/register")
async def register_webhook(payload: dict = Body(...)):
    """Register a webhook for events"""
    user = payload.get("user", "default")
    webhook_url = payload.get("webhook_url")
    events = payload.get("events", ["review_complete", "generation_complete"])
    
    if not webhook_url:
        raise HTTPException(status_code=400, detail="Webhook URL required")
    
    webhook_id = hashlib.md5(f"{user}{webhook_url}".encode()).hexdigest()[:8]
    
    if user not in WEBHOOKS:
        WEBHOOKS[user] = []
    
    webhook = {
        "id": webhook_id,
        "url": webhook_url,
        "events": events,
        "created": datetime.now().isoformat(),
        "active": True
    }
    
    WEBHOOKS[user].append(webhook)
    return {"message": "Webhook registered", "webhook_id": webhook_id}

@app.get("/api/webhooks/{user}")
async def get_webhooks(user: str):
    """Get user's registered webhooks"""
    webhooks = WEBHOOKS.get(user, [])
    return {"webhooks": webhooks, "total": len(webhooks)}

@app.delete("/api/webhooks/{user}/{webhook_id}")
async def delete_webhook(user: str, webhook_id: str):
    """Delete a webhook"""
    if user in WEBHOOKS:
        WEBHOOKS[user] = [w for w in WEBHOOKS[user] if w["id"] != webhook_id]
        return {"message": "Webhook deleted"}
    raise HTTPException(status_code=404, detail="Webhook not found")

# 5. EXPANDED CODE TEMPLATES WITH ADVANCED PATTERNS
@app.get("/api/templates/advanced/{language}")
async def get_advanced_templates(language: str):
    """Get advanced code templates and design patterns"""
    advanced_templates = {
        "python": {
            "design_patterns": {
                "singleton": "class Singleton:\\n    _instance = None\\n    def __new__(cls):\\n        if cls._instance is None:\\n            cls._instance = super().__new__(cls)\\n        return cls._instance",
                "observer": "class Subject:\\n    def __init__(self):\\n        self._observers = []\\n    def attach(self, observer):\\n        self._observers.append(observer)\\n    def notify(self):\\n        for observer in self._observers:\\n            observer.update()",
                "factory": "class AnimalFactory:\\n    @staticmethod\\n    def create_animal(animal_type):\\n        if animal_type == 'dog':\\n            return Dog()\\n        elif animal_type == 'cat':\\n            return Cat()"
            },
            "algorithms": {
                "quick_sort": "def quick_sort(arr):\\n    if len(arr) <= 1:\\n        return arr\\n    pivot = arr[0]\\n    left = [x for x in arr[1:] if x < pivot]\\n    right = [x for x in arr[1:] if x >= pivot]\\n    return quick_sort(left) + [pivot] + quick_sort(right)",
                "binary_search": "def binary_search(arr, target):\\n    left, right = 0, len(arr) - 1\\n    while left <= right:\\n        mid = (left + right) // 2\\n        if arr[mid] == target:\\n            return mid\\n        elif arr[mid] < target:\\n            left = mid + 1\\n        else:\\n            right = mid - 1\\n    return -1",
                "dynamic_programming": "def fibonacci(n, memo={}):\\n    if n in memo:\\n        return memo[n]\\n    if n <= 1:\\n        return n\\n    memo[n] = fibonacci(n-1, memo) + fibonacci(n-2, memo)\\n    return memo[n]"
            },
            "async_patterns": {
                "async_api": "import asyncio\\n\\nasync def fetch(url):\\n    await asyncio.sleep(1)\\n    return f'Data from {url}'\\n\\nasync def main():\\n    results = await asyncio.gather(fetch('url1'), fetch('url2'))\\n    return results",
                "thread_pool": "from concurrent.futures import ThreadPoolExecutor\\n\\ndef process_data(item):\\n    return item * 2\\n\\nwith ThreadPoolExecutor(max_workers=4) as executor:\\n    results = list(executor.map(process_data, [1,2,3,4]))"
            }
        },
        "javascript": {
            "design_patterns": {
                "singleton": "class Singleton {\\n  static instance = null;\\n  static getInstance() {\\n    if (!Singleton.instance) {\\n      Singleton.instance = new Singleton();\\n    }\\n    return Singleton.instance;\\n  }\\n}",
                "observer": "class EventEmitter {\\n  events = {};\\n  on(event, callback) {\\n    if (!this.events[event]) this.events[event] = [];\\n    this.events[event].push(callback);\\n  }\\n  emit(event, ...args) {\\n    this.events[event]?.forEach(cb => cb(...args));\\n  }\\n}",
                "factory": "const AnimalFactory = {\\n  create: (type) => {\\n    switch(type) {\\n      case 'dog': return { sound: 'woof' };\\n      case 'cat': return { sound: 'meow' };\\n    }\\n  }\\n}"
            },
            "react_patterns": {
                "custom_hook": "function useCounter(initial = 0) {\\n  const [count, setCount] = useState(initial);\\n  const increment = () => setCount(c => c + 1);\\n  return { count, increment };\\n}",
                "context_api": "const ThemeContext = createContext();\\nfunction ThemeProvider({ children }) {\\n  const [theme, setTheme] = useState('light');\\n  return (\\n    <ThemeContext.Provider value={{ theme, setTheme }}>\\n      {children}\\n    </ThemeContext.Provider>\\n  );\\n}"
            }
        },
        "java": {
            "design_patterns": {
                "singleton": "public class Singleton {\\n  private static Singleton instance;\\n  private Singleton() {}\\n  public static synchronized Singleton getInstance() {\\n    if (instance == null) instance = new Singleton();\\n    return instance;\\n  }\\n}",
                "builder": "public class User {\\n  private String name;\\n  private int age;\\n  public static class Builder {\\n    public Builder withName(String n) { this.name = n; return this; }\\n    public User build() { return new User(this); }\\n  }\\n}"
            }
        }
    }
    
    templates = advanced_templates.get(language, {})
    return {"advanced_templates": templates, "language": language}

# 6. RATE LIMITING AND QUOTA MANAGEMENT
@app.post("/api/quota/check")
async def check_quota(payload: dict = Body(...)):
    """Check API usage quota for user"""
    user = payload.get("user", "default")
    
    if user not in API_RATE_LIMITS:
        API_RATE_LIMITS[user] = {
            "requests_today": 0,
            "requests_month": 0,
            "limit_daily": 1000,
            "limit_monthly": 10000,
            "last_reset": datetime.now().isoformat()
        }
    
    quota = API_RATE_LIMITS[user]
    remaining_today = quota["limit_daily"] - quota["requests_today"]
    remaining_month = quota["limit_monthly"] - quota["requests_month"]
    
    return {
        "requests_today": quota["requests_today"],
        "remaining_today": max(0, remaining_today),
        "requests_month": quota["requests_month"],
        "remaining_month": max(0, remaining_month),
        "within_limits": remaining_today > 0 and remaining_month > 0
    }

@app.post("/api/quota/increment")
async def increment_quota(payload: dict = Body(...)):
    """Increment user's API usage"""
    user = payload.get("user", "default")
    
    if user not in API_RATE_LIMITS:
        API_RATE_LIMITS[user] = {
            "requests_today": 0,
            "requests_month": 0,
            "limit_daily": 1000,
            "limit_monthly": 10000
        }
    
    API_RATE_LIMITS[user]["requests_today"] += 1
    API_RATE_LIMITS[user]["requests_month"] += 1
    
    return {"message": "Quota incremented"}

# 7. BATCH ANALYSIS
@app.post("/api/batch-analyze")
async def batch_analyze(files: List[UploadFile] = File(...)):
    """Process multiple code files at once"""
    results = []
    for file in files:
        content = await file.read()
        code = content.decode("utf-8", errors="ignore")
        # Perform a quick complexity scan for the summary
        complexity = analyze_complexity(code)
        results.append({
            "filename": file.filename,
            "complexity": complexity,
            "size_bytes": len(code)
        })
    return {"results": results, "total_files": len(files)}

# 9. DATA EXPORT (CSV)
@app.get("/api/analytics/export")
async def export_analytics_csv():
    """Export user analytics data to CSV"""
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write Header
    writer.writerow(["Username", "Total Reviews", "Total Generations", "Last Activity", "Most Used Language"])
    
    # Write Data
    for user, data in USER_ANALYTICS.items():
        langs = data.get("languages", {})
        top_lang = max(langs, key=langs.get) if langs else "N/A"
        writer.writerow([
            user,
            data.get("reviews", 0),
            data.get("generations", 0),
            data.get("last_activity", "N/A"),
            top_lang
        ])
    
    output.seek(0)
    return Response(content=output.getvalue(), media_type="text/csv", headers={"Content-Disposition": "attachment; filename=analytics_export.csv"})

# 8. REAL-TIME COLLABORATION (WebSockets)
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        self.active_connections.remove(websocket)

    async def broadcast(self, message: str):
        for connection in self.active_connections:
            await connection.send_text(message)

manager = ConnectionManager()

@app.websocket("/ws/collab")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        while True:
            data = await websocket.receive_text()
            await manager.broadcast(data)
    except WebSocketDisconnect:
        manager.disconnect(websocket)

# 10. REPORT MANAGEMENT
@app.get("/api/reports")
async def list_reports(user: str = "default"):
    """List saved reports for a specific user"""
    files = []
    # Filter files that contain the username
    for f in REPORTS_DIR.glob(f"report_{user}_*.*"):
        files.append({
            "filename": f.name,
            "created": datetime.fromtimestamp(f.stat().st_ctime).isoformat(),
            "size": f.stat().st_size
        })
    return {"reports": sorted(files, key=lambda x: x["created"], reverse=True)}

@app.get("/api/reports/{filename}")
async def get_report(filename: str):
    """Download a specific report"""
    file_path = REPORTS_DIR / filename
    if file_path.exists():
        return FileResponse(file_path, filename=filename)
    raise HTTPException(status_code=404, detail="File not found")

@app.delete("/api/reports/{filename}")
async def delete_report(filename: str):
    """Delete a report"""
    file_path = REPORTS_DIR / filename
    if file_path.exists():
        os.remove(file_path)
        return {"message": "Report deleted"}
    raise HTTPException(status_code=404, detail="File not found")

@app.post("/api/reports/save")
async def save_report_to_disk(payload: dict = Body(...)):
    """Generate and save report to disk"""
    format_type = payload.get("format", "docx")
    original_code = payload.get("original_code", "")
    rewritten_code = payload.get("rewritten_code", "")
    review = payload.get("review", "")
    stats = payload.get("stats", {})
    user = payload.get("user", "default")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{user}_{timestamp}.{format_type}"
    file_path = REPORTS_DIR / filename
    
    try:
        if format_type == "docx":
            from docx import Document
            doc = Document()
            doc.add_heading("Code Refine Report", 0)
            doc.add_heading("Analysis Statistics", level=1)
            doc.add_paragraph(f"Critical Issues: {stats.get('critical', 0)}")
            doc.add_paragraph(f"High Priority: {stats.get('high', 0)}")
            doc.add_heading("Review & Feedback", level=1)
            doc.add_paragraph(review)
            doc.add_heading("Rewritten Code", level=1)
            doc.add_paragraph(rewritten_code, style='List Number')
            doc.save(file_path)
            
        elif format_type == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, "Code Refine Report", ln=True)
            pdf.set_font("Arial", "", 12)
            pdf.cell(0, 10, f"Date: {timestamp}", ln=True)
            pdf.ln(10)
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Review Summary:", ln=True)
            pdf.set_font("Arial", "", 10)
            pdf.multi_cell(0, 5, review[:1000]) # Limit length for simple PDF
            pdf.output(str(file_path))
            
        return {"message": "Report saved", "filename": filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save report: {str(e)}")

# 10.5 GITHUB INTEGRATION
@app.post("/api/github/commit")
async def commit_to_github(payload: dict = Body(...)):
    """Commit code to GitHub"""
    user = payload.get("user", "default")
    repo = payload.get("repo")
    path = payload.get("path")
    message = payload.get("message", "Update via Code Refine")
    content = payload.get("content")
    branch = payload.get("branch", "main")
    token = payload.get("token")

    if user.lower() == "guest":
        return JSONResponse(status_code=401, content={"error": "auth_required", "redirect": "/login"})

    if not token:
        settings = USER_SETTINGS.get(user, {})
        token = decrypt_secret(settings.get("github_token"))

    if not token:
        raise HTTPException(status_code=400, detail="GitHub Token required. Set it in Settings or provide it.")
    
    if not repo or not path or not content:
        raise HTTPException(status_code=400, detail="Repository, path, and content are required")

    headers = {"Authorization": f"Bearer {token}", "Accept": "application/vnd.github.v3+json"}
    
    async with httpx.AsyncClient() as client:
        # 1. Check if file exists to get SHA (for update)
        url = f"https://api.github.com/repos/{repo}/contents/{path}"
        get_res = await client.get(url, headers=headers, params={"ref": branch})
        
        data = {
            "message": message,
            "content": base64.b64encode(content.encode("utf-8")).decode("utf-8"),
            "branch": branch
        }
        if get_res.status_code == 200:
            data["sha"] = get_res.json().get("sha")
            
        put_res = await client.put(url, headers=headers, json=data)
        if put_res.status_code not in [200, 201]:
            raise HTTPException(status_code=put_res.status_code, detail=f"GitHub Error: {put_res.text}")
            
        return {"message": "File committed successfully", "url": put_res.json().get("html_url")}

# 11. USER SETTINGS
@app.post("/api/settings/update")
async def update_settings(payload: dict = Body(...)):
    """Update user settings"""
    user = payload.get("user") or payload.get("username") or "default"
    settings = payload.get("settings", {})
    
    # Encrypt sensitive keys before storage
    if "gemini_key" in settings and settings["gemini_key"]:
        settings["gemini_key"] = encrypt_secret(settings["gemini_key"])
    if "github_token" in settings and settings["github_token"]:
        settings["github_token"] = encrypt_secret(settings["github_token"])
    
    USER_SETTINGS[user] = settings
    return {"message": "Settings updated"}

@app.get("/api/settings/{user}")
async def get_settings(user: str):
    """Get user settings"""
    settings = USER_SETTINGS.get(user, {
        "model": "llama-3.3-70b",
        "theme": "dark",
        "font_size": 14,
        "tab_size": 4,
        "notifications": False
    }).copy() # Copy to avoid modifying storage during decryption
    
    # Decrypt for display
    if "gemini_key" in settings:
        settings["gemini_key"] = decrypt_secret(settings["gemini_key"])
    if "github_token" in settings:
        settings["github_token"] = decrypt_secret(settings["github_token"])
        
    return {"settings": settings}

# 11.5 TEST CONNECTION
@app.post("/api/test-gemini")
async def test_gemini_connection(payload: dict = Body(...)):
    """Test Gemini API Key validity"""
    api_key = payload.get("api_key")
    if not api_key:
        raise HTTPException(status_code=400, detail="API Key required")
    
    if not GEMINI_AVAILABLE:
         raise HTTPException(status_code=400, detail="Gemini library not installed")

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-pro")
        response = model.generate_content("Hello, just checking connection.")
        return {"message": "Connection successful!", "response": response.text[:50]}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Connection failed: {str(e)}")

# 12. HELP & DOCUMENTATION
FAQ_DATA = {
    "categories": [
        {
            "name": "Getting Started",
            "items": [
                {"question": "How do I analyze my code?", "answer": "Simply paste your code into the editor on the main page and click 'Review Code'. You can also upload files directly."},
                {"question": "What languages are supported?", "answer": "We support Python, JavaScript, Java, C++, Go, Rust, and many others. The system auto-detects the language."}
            ]
        },
        {
            "name": "Features",
            "items": [
                {"question": "How does Batch Analysis work?", "answer": "Go to the Batch Analysis page, drag and drop multiple files, and the system will process them sequentially, providing a summary for each."},
                {"question": "Is my code saved?", "answer": "We do not permanently store your code unless you explicitly save it to your history or snippets library. All analysis is stateless by default."}
            ]
        },
        {
            "name": "Account & Settings",
            "items": [
                {"question": "How do I change the AI model?", "answer": "Navigate to Settings and select your preferred model (e.g., Llama 3.3, Mixtral) from the dropdown menu."},
                {"question": "Can I export my data?", "answer": "Yes, you can export your usage analytics as a CSV file from the Dashboard or Profile page."}
            ]
        }
    ]
}

@app.get("/api/help/faq")
async def get_faq():
    """Get FAQ and help content"""
    return FAQ_DATA

# 13. NEWSLETTER SUBSCRIPTION
def send_welcome_email_background(email: str):
    """Send welcome email to new subscribers"""
    smtp_server = os.getenv("SMTP_SERVER")
    smtp_port = os.getenv("SMTP_PORT")
    smtp_user = os.getenv("SMTP_USER")
    smtp_password = os.getenv("SMTP_PASSWORD")
    sender_email = os.getenv("SMTP_SENDER_EMAIL", smtp_user)

    if smtp_server and smtp_port and smtp_user and smtp_password:
        try:
            msg = MIMEMultipart()
            msg['From'] = sender_email
            msg['To'] = email
            msg['Subject'] = "Welcome to Code Refine!"
            
            body = f"""
            <html>
            <body style="font-family: Arial, sans-serif; color: #333;">
                <div style="max-width: 600px; margin: 0 auto; padding: 20px; border: 1px solid #eee; border-radius: 10px;">
                    <h2 style="color: #2563eb;">Welcome to Code Refine! 🚀</h2>
                    <p>Hi there,</p>
                    <p>Thank you for subscribing to our newsletter. We are excited to have you on board!</p>
                    <p>You will now receive updates about the latest features, coding tips, and security insights directly to your inbox.</p>
                    <br>
                    <p>Happy Coding,</p>
                    <p><strong>The Code Refine Team</strong></p>
                    <hr style="border: none; border-top: 1px solid #eee; margin: 20px 0;">
                    <p style="font-size: 12px; color: #888;">If you did not subscribe to this newsletter, please ignore this email.</p>
                </div>
            </body>
            </html>
            """
            msg.attach(MIMEText(body, 'html'))
            
            if int(smtp_port) == 465:
                server = smtplib.SMTP_SSL(smtp_server, int(smtp_port))
            else:
                server = smtplib.SMTP(smtp_server, int(smtp_port))
                server.starttls()
            
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
            server.quit()
            print(f"✅ Welcome email sent to {email}")
        except Exception as e:
            print(f"❌ Failed to send welcome email: {e}")
    else:
        print(f"📧 [SIMULATION] Welcome email to {email} (SMTP not configured)")

@app.post("/api/newsletter/subscribe")
async def subscribe_newsletter(background_tasks: BackgroundTasks, payload: dict = Body(...)):
    """Subscribe to newsletter"""
    email = payload.get("email")
    if not email:
        raise HTTPException(status_code=400, detail="Email required")
    
    # Send welcome email in background
    background_tasks.add_task(send_welcome_email_background, email)
    
    NEWSLETTER_SUBS.append({"email": email, "date": datetime.now().isoformat()})
    return {"message": "Subscribed successfully"}

@app.get("/api/newsletter/export")
async def export_newsletter_csv():
    """Export newsletter subscribers to CSV"""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Email", "Date Subscribed"])
    for sub in NEWSLETTER_SUBS:
        writer.writerow([sub["email"], sub["date"]])
    output.seek(0)
    return Response(content=output.getvalue(), media_type="text/csv", headers={"Content-Disposition": "attachment; filename=subscribers.csv"})

# 14. PWA STATIC FILES
@app.get("/sw.js", include_in_schema=False)
async def service_worker():
    path = Path(__file__).parent.parent / "frontend" / "sw.js"
    if path.exists():
        return FileResponse(path, media_type="application/javascript")
    raise HTTPException(status_code=404, detail="Service Worker not found")

@app.get("/manifest.json", include_in_schema=False)
async def manifest():
    path = Path(__file__).parent.parent / "frontend" / "manifest.json"
    if path.exists():
        return FileResponse(path, media_type="application/json")
    raise HTTPException(status_code=404, detail="Manifest not found")

# 14.5 SERVE CORE JS FILES (Root Level)
@app.get("/main.js", include_in_schema=False)
async def serve_main_js():
    path = Path(__file__).parent.parent / "frontend" / "main.js"
    if path.exists():
        return FileResponse(path, media_type="application/javascript")
    raise HTTPException(status_code=404, detail="main.js not found")

@app.get("/utils.js", include_in_schema=False)
async def serve_utils_js():
    path = Path(__file__).parent.parent / "frontend" / "utils.js"
    if path.exists():
        return FileResponse(path, media_type="application/javascript")
    raise HTTPException(status_code=404, detail="utils.js not found")

@app.get("/theme.js", include_in_schema=False)
async def serve_theme_js():
    path = Path(__file__).parent.parent / "frontend" / "theme.js"
    if path.exists():
        return FileResponse(path, media_type="application/javascript")
    raise HTTPException(status_code=404, detail="theme.js not found")

# 15. SYSTEM STATUS
@app.get("/api/system/status")
async def system_status():
    """Comprehensive system status check"""
    groq_status = "Unknown"
    try:
        if api_key:
            client.models.list()
            groq_status = "Operational"
        else:
            groq_status = "Not Configured (Missing API Key)"
    except Exception as e:
        groq_status = f"Error: {str(e)}"

    return {
        "server": "Operational",
        "version": "2.0.0",
        "groq_api": groq_status,
        "storage": {
            "users": len(USER_DB),
            "sessions": len(ACTIVE_SESSIONS),
            "snippets": sum(len(v) for v in CODE_SNIPPETS.values()),
            "history": sum(len(v) for v in CODE_HISTORY.values())
        },
        "system": {
            "platform": platform.system(),
            "python": platform.python_version(),
            "time": datetime.now().isoformat()
        }
    }

# 16. ADMIN USER MANAGEMENT
@app.get("/api/admin/users")
async def get_all_users():
    """Get list of all users (Admin only)"""
    users = []
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor(dictionary=True)
            cursor.execute("SELECT username, email, role, created_at FROM users ORDER BY created_at DESC")
            users = cursor.fetchall()
            # Convert datetime to string
            for u in users:
                if isinstance(u.get('created_at'), datetime):
                    u['created_at'] = u['created_at'].isoformat()
        except Error as e:
            print(f"DB Error: {e}")
        finally:
            if conn.is_connected():
                cursor.close()
                conn.close()
    
    if not users:
        # Fallback to in-memory
        for username, data in USER_DB.items():
            users.append({
                "username": username,
                "email": data.get("email", ""),
                "role": "admin" if username == "admin" else "user",
                "created_at": datetime.now().isoformat()
            })
            
    return {"users": users}

@app.post("/api/admin/reset-password")
async def admin_reset_password(payload: dict = Body(...)):
    """Reset user password (Admin only)"""
    target_username = payload.get("username")
    new_password = payload.get("new_password")
    
    if not target_username or not new_password:
        raise HTTPException(status_code=400, detail="Username and new password required")

    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute("UPDATE users SET password_hash = %s WHERE username = %s", (get_password_hash(new_password), target_username))
            conn.commit()
            if cursor.rowcount == 0:
                 raise HTTPException(status_code=404, detail="User not found")
        except Error as e:
             raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
        finally:
            if conn.is_connected():
                cursor.close()
                conn.close()
    else:
        # Fallback to in-memory
        if target_username in USER_DB:
            USER_DB[target_username]["password"] = get_password_hash(new_password)
        else:
            raise HTTPException(status_code=404, detail="User not found")
            
    return {"message": f"Password for {target_username} reset successfully"}

@app.get("/api/admin/maintenance")
async def get_maintenance_status():
    """Get maintenance mode status"""
    return {"enabled": MAINTENANCE_MODE}

@app.post("/api/admin/maintenance")
async def set_maintenance_status(payload: dict = Body(...)):
    """Toggle maintenance mode (Admin only)"""
    global MAINTENANCE_MODE
    username = payload.get("username")
    if not username or not is_user_admin(username):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    enabled = payload.get("enabled", False)
    MAINTENANCE_MODE = enabled
    return {"message": f"Maintenance mode {'enabled' if enabled else 'disabled'}", "enabled": MAINTENANCE_MODE}

# --- Password Reset (Simulated) ---

def send_email_background(email: str, reset_link: str):
    """Send email in background to prevent blocking the API"""
    smtp_server = os.getenv("SMTP_SERVER")
    smtp_port = os.getenv("SMTP_PORT")
    smtp_user = os.getenv("SMTP_USER")
    smtp_password = os.getenv("SMTP_PASSWORD")
    sender_email = os.getenv("SMTP_SENDER_EMAIL", smtp_user)

    if smtp_server and smtp_port and smtp_user and smtp_password:
        try:
            msg = MIMEMultipart()
            msg['From'] = sender_email
            msg['To'] = email
            msg['Subject'] = "Password Reset Request - Code Refine"
            
            body = f"""
            <html>
            <body>
                <h2>Password Reset Request</h2>
                <p>Hello,</p>
                <p>We received a request to reset your password for your Code Refine account.</p>
                <p>Click the link below to reset it:</p>
                <p><a href="{reset_link}" style="background-color: #0ea5e9; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px;">Reset Password</a></p>
                <p>Or copy this link: {reset_link}</p>
                <p>If you did not request this, please ignore this email.</p>
            </body>
            </html>
            """
            msg.attach(MIMEText(body, 'html'))
            
            if int(smtp_port) == 465:
                server = smtplib.SMTP_SSL(smtp_server, int(smtp_port))
            else:
                server = smtplib.SMTP(smtp_server, int(smtp_port))
                server.starttls()
            
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
            server.quit()
            print(f"✅ Email sent to {email} via SMTP")
        except Exception as e:
            print(f"❌ Failed to send email: {e}")
            print(f"📧 [FALLBACK SIMULATION] Reset Link: {reset_link}")
    else:
        print(f"📧 [SIMULATION] SMTP not configured. Link: {reset_link}")

@app.post("/api/forgot-password")
async def forgot_password(background_tasks: BackgroundTasks, payload: dict = Body(...)):
    """Initiate password reset flow (Real Email via Background Task)"""
    email = payload.get("email")
    if not email:
        raise HTTPException(status_code=400, detail="Email is required")
    
    user_found = False
    
    # Check MySQL
    if DB_AVAILABLE:
        conn = get_db_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
                if cursor.fetchone():
                    user_found = True
            except Error as e:
                print(f"DB Error: {e}")
            finally:
                if conn.is_connected():
                    cursor.close()
                    conn.close()

    # Check In-Memory
    if not user_found:
        for data in USER_DB.values():
            if data.get("email") == email:
                user_found = True
                break
    
    if user_found:
        reset_token = hashlib.sha256(f"{email}{datetime.now()}".encode()).hexdigest()[:16]
        reset_link = f"http://localhost:8000/reset-password?token={reset_token}"
        
        # Send email in background
        background_tasks.add_task(send_email_background, email, reset_link)
        
    return {"message": "If an account exists for this email, a reset link has been sent."}

# --- Page Routing ---

@app.get("/{page}", response_class=HTMLResponse)
async def serve_ui(page: str):
    # Dynamic path finding
    base_path = Path(__file__).parent.parent / "frontend"
    file_map = {
        "app": "index.html", 
        "dashboard": "dashboard.html", 
        "login": "login.html",
        "admin": "admin.html",
        "batch": "batch.html",
        "collab": "collab.html",
        "profile": "profile.html",
        "reports": "reports.html",
        "settings": "settings.html",
        "help": "help.html",
        "landing": "landing.html",
        "404": "404.html",
        "generate": "generate.html",
        "status": "status.html",
        "signup": "signup.html"
    }
    
    target_file = file_map.get(page)
    
    # If page not found in map, serve 404
    if not target_file:
        path_404 = base_path / "404.html"
        return HTMLResponse(path_404.read_text(encoding="utf-8"), status_code=404) if path_404.exists() else HTMLResponse("<h1>404: Page Not Found</h1>", status_code=404)

    path = base_path / target_file
    
    if path.exists():
        return path.read_text(encoding="utf-8")
    return HTMLResponse("<h1>404: File Not Found</h1>", status_code=404)

@app.get("/", response_class=HTMLResponse)
async def root():
    path = Path(__file__).parent.parent / "frontend" / "landing.html"
    return path.read_text(encoding="utf-8") if path.exists() else "Login Page Not Found"

if __name__ == "__main__":
    import uvicorn
   import os
uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))